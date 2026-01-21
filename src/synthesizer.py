from typing import List
from deepeval.synthesizer import Synthesizer
from deepeval.synthesizer.config import EvolutionConfig, Evolution, StylingConfig, FiltrationConfig

import os
import json
from pathlib import Path

class DatasetGenerator:
    def __init__(self, config: dict, base_dir: Path = None):
        # Base directory for this session
        self.base_dir = base_dir if base_dir else Path("data")
        self.source_docs_dir = self.base_dir / "source_docs"
        self.synthetic_data_dir = self.base_dir / "synthetic_data"
        
        # Ensure directories exist
        self.source_docs_dir.mkdir(parents=True, exist_ok=True)
        self.synthetic_data_dir.mkdir(parents=True, exist_ok=True)

        # Extract settings from the passed config
        api_key = config.get("api_key") or os.getenv("OPENAI_API_KEY")
        model_name = config.get("model_name", "gpt-4o-mini")
        
        if not api_key:
            raise ValueError("OpenAI API key is empty. Please configure a valid key.")
        
        # Initialize model directly with config values
        from deepeval.models import GPTModel
        self.model = GPTModel(model=model_name, api_key=api_key)

        # Config defaults
        reasoning_weight = config.get("reasoning_weight", 0.5)
        multicontext_weight = config.get("multicontext_weight", 0.5)
        task = config.get("task", "Expert Customer Support")
        scenario = config.get("scenario", "A customer interacting with an automated assistant.")
        input_format = config.get("input_format", "Professional and specific queries")
        expected_output_format = config.get("expected_output_format", "Detailed responses with citations")
        
        # Advanced settings
        self.num_evolutions = config.get("num_evolutions", 2)
        self.num_goldens = config.get("num_goldens", 2)

        self.evolution_config = EvolutionConfig(
            evolutions={
                Evolution.REASONING: reasoning_weight,
                Evolution.MULTICONTEXT: multicontext_weight
            },
            num_evolutions=self.num_evolutions
        )
        self.styling_config = StylingConfig(
            task=task,
            scenario=scenario,
            input_format=input_format,
            expected_output_format=expected_output_format
        )
        self.synthesizer = Synthesizer(
            model=self.model,
            evolution_config=self.evolution_config,
            styling_config=self.styling_config,
            filtration_config=FiltrationConfig()
        )
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using pdfplumber (Windows-compatible)."""
        import pdfplumber
        text_parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    
    def _extract_from_excel(self, path: str) -> str:
        import pandas as pd
        try:
            df = pd.read_excel(path)
            # Convert to string, handling NaNs
            return df.to_string(index=False, na_rep="")
        except Exception as e:
            return f"Error reading Excel {path}: {str(e)}"

    def _extract_from_csv(self, path: str) -> str:
        import pandas as pd
        try:
            df = pd.read_csv(path)
            return df.to_string(index=False, na_rep="")
        except Exception as e:
            return f"Error reading CSV {path}: {str(e)}"

    def _extract_from_txt(self, path: str) -> str:
        try:
            return Path(path).read_text(encoding='utf-8', errors='replace')
        except Exception as e:
            return f"Error reading TXT {path}: {str(e)}"

    def _get_document_contexts(self, paths: List[str]) -> List[str]:
        """Extract text contexts from documents, handling all supported formats."""
        contexts = []
        for path in paths:
            path_obj = Path(path)
            text = ""
            
            try:
                suffix = path_obj.suffix.lower()
                if suffix == '.pdf':
                    text = self._extract_text_from_pdf(str(path))
                elif suffix == '.docx':
                    from docx import Document
                    doc = Document(path)
                    
                    # Extract paragraphs
                    full_text = [para.text for para in doc.paragraphs if para.text.strip()]
                    
                    # Extract tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                full_text.append(" | ".join(row_text))
                                
                    text = "\n".join(full_text)
                elif suffix == '.xlsx':
                    text = self._extract_from_excel(str(path))
                elif suffix == '.csv':
                    text = self._extract_from_csv(str(path))
                elif suffix == '.txt':
                    text = self._extract_from_txt(str(path))
                
                if text.strip():
                    # Split into chunks of ~2000 chars for better context handling
                    chunk_size = 2000
                    for i in range(0, len(text), chunk_size):
                        chunk = text[i:i+chunk_size].strip()
                        if chunk:
                            contexts.append(chunk)
                            
            except Exception:
                pass
                
        return contexts

    def generate_single_turn(self, paths: List[str]):
        self.synthesizer.synthetic_goldens = []
        all_goldens = []
        last_error = None
        
        for path in paths:
            file_name = Path(path).name
            contexts = self._get_document_contexts([path])
            
            if not contexts:
                print(f"Warning: No text extracted from {file_name}")
                continue

            # Contexts must be List[List[str]]
            formatted_contexts = [[c] for c in contexts]

            # Generate for this specific file
            try:
                # We use a temporary list to capture just this batch
                batch_goldens = self.synthesizer.generate_goldens_from_contexts(
                    contexts=formatted_contexts,
                    max_goldens_per_context=self.num_goldens,
                    include_expected_output=True
                )
                
                # Manually inject source file
                for golden in batch_goldens:
                    golden.source_file = file_name
                    
                all_goldens.extend(batch_goldens)
                
            except Exception as e:
                print(f"Error generating for {file_name}: {e}")
                last_error = e

        if not all_goldens:
            # Re-raise the actual error if we have one, otherwise generic
            if last_error:
                raise last_error
            raise ValueError("No goldens were generated. Please check your documents or API key.")
        
        # Update the synthesizer's list so save_as works correctly
        self.synthesizer.synthetic_goldens = all_goldens
        self.synthesizer.save_as(file_type='json', directory=str(self.synthetic_data_dir), file_name="single_turn_goldens")
        return all_goldens

    def generate_multi_turn(self, paths: List[str]):
        self.synthesizer.synthetic_goldens = []
        self.synthesizer.synthetic_conversational_goldens = []
        all_conversations = []
        last_error = None
        
        for path in paths:
            file_name = Path(path).name
            contexts = self._get_document_contexts([path])
            
            if not contexts:
                continue
                
            # Contexts must be List[List[str]]
            formatted_contexts = [[c] for c in contexts]
                
            try:
                batch_conversations = []
                 # Try new method name first (following v3 naming convention)
                if hasattr(self.synthesizer, 'generate_conversational_goldens_from_contexts'):
                     batch_conversations = self.synthesizer.generate_conversational_goldens_from_contexts(
                        contexts=formatted_contexts,
                        max_goldens_per_context=self.num_goldens
                     )
                else:
                    # Fallback to legacy method name
                    batch_conversations = self.synthesizer.generate_conversational_goldens(
                        contexts=formatted_contexts,
                        max_goldens_per_context=self.num_goldens
                    )
                
                # Manually inject source file (multi-turn structure might differ, checking standard property)
                for conv in batch_conversations:
                    if hasattr(conv, 'source_file'):
                        conv.source_file = file_name
                    # Standardize expected_output for frontend consistency
                    if hasattr(conv, 'expected_outcome'):
                        conv.expected_output = conv.expected_outcome
                        
                all_conversations.extend(batch_conversations)

            except Exception as e:
                last_error = e
                # If one method fails, try the other as fallback (legacy behavior fallback logic preserved)
                try:
                     batch_conversations = self.synthesizer.generate_conversational_goldens(
                        contexts=formatted_contexts,
                        max_goldens_per_context=self.num_goldens
                     )
                     for conv in batch_conversations:
                        if hasattr(conv, 'source_file'):
                            conv.source_file = file_name
                        # Standardize expected_output for frontend consistency
                        if hasattr(conv, 'expected_outcome'):
                            conv.expected_output = conv.expected_outcome
                            
                     all_conversations.extend(batch_conversations)
                     last_error = None # Clear error if fallback worked
                except Exception:
                     print(f"Error generating multi-turn for {file_name}: {e}")
        
        if not all_conversations:
            if last_error:
                raise last_error
            raise ValueError("No conversations were generated. Please check your documents.")
        
        # Update proper list for saving
        self.synthesizer.synthetic_conversational_goldens = all_conversations
        self.synthesizer.save_as(file_type='json', directory=str(self.synthetic_data_dir), file_name="multi_turn_goldens")
        return all_conversations